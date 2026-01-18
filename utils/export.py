"""
Export utilities - Convert appendix content to richly formatted DOCX and PDF files.

Supports:
- Headings (levels 1-5)
- Bold, italic, underline
- Superscript, subscript
- Formatted tables with borders and shading
- Mathematical symbols and equations
- Bullet and numbered lists
- Blockquotes
- Code blocks
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, ListFlowable, ListItem
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
import io
import re


# ============================================================================
# MARKDOWN PARSING UTILITIES
# ============================================================================

def parse_markdown_content(content: str) -> list:
    """
    Parse markdown content into structured elements.
    Returns a list of dictionaries with type and content.
    """
    elements = []
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Empty line
        if not line.strip():
            elements.append({'type': 'empty'})
            i += 1
            continue
        
        # Headers (levels 1-5)
        header_match = re.match(r'^(#{1,5})\s+(.+)$', line)
        if header_match:
            level = len(header_match.group(1))
            text = header_match.group(2)
            elements.append({'type': f'heading{level}', 'text': text})
            i += 1
            continue
        
        # Table detection
        if '|' in line and i + 1 < len(lines) and re.match(r'^[\|\s\-:]+$', lines[i + 1]):
            table_lines = [line]
            i += 1
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            elements.append({'type': 'table', 'lines': table_lines})
            continue
        
        # Code block
        if line.strip().startswith('```'):
            code_lines = []
            language = line.strip()[3:]
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # Skip closing ```
            elements.append({'type': 'code', 'language': language, 'content': '\n'.join(code_lines)})
            continue
        
        # Blockquote
        if line.strip().startswith('>'):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                quote_lines.append(lines[i].strip()[1:].strip())
                i += 1
            elements.append({'type': 'blockquote', 'text': ' '.join(quote_lines)})
            continue
        
        # Bullet list
        if re.match(r'^[\s]*[-*]\s+', line):
            list_items = []
            while i < len(lines) and re.match(r'^[\s]*[-*]\s+', lines[i]):
                item_text = re.sub(r'^[\s]*[-*]\s+', '', lines[i])
                list_items.append(item_text)
                i += 1
            elements.append({'type': 'bullet_list', 'items': list_items})
            continue
        
        # Numbered list
        if re.match(r'^[\s]*\d+\.\s+', line):
            list_items = []
            while i < len(lines) and re.match(r'^[\s]*\d+\.\s+', lines[i]):
                item_text = re.sub(r'^[\s]*\d+\.\s+', '', lines[i])
                list_items.append(item_text)
                i += 1
            elements.append({'type': 'numbered_list', 'items': list_items})
            continue
        
        # Regular paragraph
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(r'^#{1,5}\s+', lines[i]) and '|' not in lines[i] and not lines[i].strip().startswith(('```', '>', '-', '*', '1.')):
            para_lines.append(lines[i])
            i += 1
        elements.append({'type': 'paragraph', 'text': ' '.join(para_lines)})
    
    return elements


def parse_inline_formatting(text: str) -> list:
    """
    Parse inline formatting (bold, italic, underline, superscript, subscript).
    Returns list of tuples: (text, formatting_dict)
    """
    result = []
    
    # Pattern for various inline formats
    # Order matters: process longer patterns first
    patterns = [
        (r'\*\*\*(.+?)\*\*\*', {'bold': True, 'italic': True}),  # Bold+Italic
        (r'\*\*(.+?)\*\*', {'bold': True}),  # Bold
        (r'__(.+?)__', {'bold': True}),  # Bold (alt)
        (r'\*(.+?)\*', {'italic': True}),  # Italic
        (r'_(.+?)_', {'italic': True}),  # Italic (alt)
        (r'~~(.+?)~~', {'strikethrough': True}),  # Strikethrough
        (r'<u>(.+?)</u>', {'underline': True}),  # Underline
        (r'<sup>(.+?)</sup>', {'superscript': True}),  # Superscript
        (r'<sub>(.+?)</sub>', {'subscript': True}),  # Subscript
        (r'\^(.+?)\^', {'superscript': True}),  # Superscript (alt)
        (r'~(.+?)~', {'subscript': True}),  # Subscript (alt)
        (r'`(.+?)`', {'code': True}),  # Inline code
    ]
    
    # Simple approach: process text sequentially
    remaining = text
    
    while remaining:
        earliest_match = None
        earliest_pos = len(remaining)
        matched_pattern = None
        matched_format = None
        
        for pattern, formatting in patterns:
            match = re.search(pattern, remaining)
            if match and match.start() < earliest_pos:
                earliest_match = match
                earliest_pos = match.start()
                matched_pattern = pattern
                matched_format = formatting
        
        if earliest_match:
            # Add text before match
            if earliest_pos > 0:
                result.append((remaining[:earliest_pos], {}))
            # Add formatted text
            result.append((earliest_match.group(1), matched_format))
            remaining = remaining[earliest_match.end():]
        else:
            # No more matches
            if remaining:
                result.append((remaining, {}))
            break
    
    return result if result else [(text, {})]


def parse_table(table_lines: list) -> dict:
    """Parse markdown table into structured data."""
    if len(table_lines) < 2:
        return {'headers': [], 'rows': []}
    
    # Parse header
    header_line = table_lines[0]
    headers = [cell.strip() for cell in header_line.split('|')[1:-1]]
    
    # Skip separator line (index 1)
    
    # Parse data rows
    rows = []
    for line in table_lines[2:]:
        if '|' in line:
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            # Pad cells if needed
            while len(cells) < len(headers):
                cells.append('')
            rows.append(cells[:len(headers)])  # Trim if too many
    
    return {'headers': headers, 'rows': rows}


# ============================================================================
# DOCX EXPORT
# ============================================================================

def setup_docx_styles(doc: Document):
    """Set up custom styles for the document."""
    styles = doc.styles
    
    # Heading styles (1-5)
    heading_sizes = {1: 24, 2: 20, 3: 16, 4: 14, 5: 12}
    heading_colors = {1: '1E3A5F', 2: '2E5A7C', 3: '3E7A9C', 4: '4E8AAC', 5: '5E9ABC'}
    
    for level in range(1, 6):
        style_name = f'Custom Heading {level}'
        try:
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        except:
            style = styles[style_name]
        
        style.font.size = Pt(heading_sizes[level])
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(heading_colors[level])
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
    
    # Blockquote style
    try:
        quote_style = styles.add_style('Blockquote', WD_STYLE_TYPE.PARAGRAPH)
    except:
        quote_style = styles['Blockquote']
    quote_style.font.italic = True
    quote_style.font.color.rgb = RGBColor(80, 80, 80)
    quote_style.paragraph_format.left_indent = Cm(1)
    
    # Code style
    try:
        code_style = styles.add_style('Code Block', WD_STYLE_TYPE.PARAGRAPH)
    except:
        code_style = styles['Code Block']
    code_style.font.name = 'Courier New'
    code_style.font.size = Pt(9)
    code_style.paragraph_format.left_indent = Cm(0.5)


def add_formatted_run(paragraph, text: str, formatting: dict):
    """Add a run with specific formatting to a paragraph."""
    run = paragraph.add_run(text)
    
    if formatting.get('bold'):
        run.bold = True
    if formatting.get('italic'):
        run.italic = True
    if formatting.get('underline'):
        run.underline = True
    if formatting.get('strikethrough'):
        run.font.strike = True
    if formatting.get('code'):
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    if formatting.get('superscript'):
        run.font.superscript = True
    if formatting.get('subscript'):
        run.font.subscript = True


def add_formatted_paragraph(doc: Document, text: str, style=None):
    """Add a paragraph with inline formatting."""
    para = doc.add_paragraph()
    if style:
        para.style = style
    
    parts = parse_inline_formatting(text)
    for part_text, formatting in parts:
        add_formatted_run(para, part_text, formatting)
    
    return para


def set_table_borders(table):
    """Set borders for all cells in a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def add_table_to_docx(doc: Document, table_data: dict):
    """Add a formatted table to the document."""
    headers = table_data['headers']
    rows = table_data['rows']
    
    if not headers:
        return
    
    num_cols = len(headers)
    num_rows = len(rows) + 1  # +1 for header
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    set_table_borders(table)
    
    # Header row
    header_row = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ''
        para = cell.paragraphs[0]
        run = para.add_run(header_text)
        run.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Header shading
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D9E2F3')
        cell._tc.get_or_add_tcPr().append(shading)
    
    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ''
            para = cell.paragraphs[0]
            parts = parse_inline_formatting(cell_text)
            for part_text, formatting in parts:
                add_formatted_run(para, part_text, formatting)
    
    doc.add_paragraph()  # Space after table


def export_to_docx(content: str, title: str) -> bytes:
    """
    Export Markdown content as a richly formatted Word document with metadata.
    """
    from datetime import datetime

    doc = Document()

    # Set document properties (metadata)
    core_properties = doc.core_properties
    core_properties.title = title
    core_properties.author = "Forward Thinking - Foresight"
    core_properties.comments = "Generated by Forward Thinking - Foresight AI Assistant"
    core_properties.created = datetime.now()

    # Set up styles
    setup_docx_styles(doc)

    # Add title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor.from_string('1E3A5F')
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.space_after = Pt(24)
    
    # Parse and add content
    elements = parse_markdown_content(content)
    
    for element in elements:
        elem_type = element.get('type')
        
        if elem_type == 'empty':
            doc.add_paragraph()
        
        elif elem_type.startswith('heading'):
            level = int(elem_type[-1])
            para = doc.add_paragraph()
            run = para.add_run(element['text'])
            run.bold = True
            sizes = {1: 24, 2: 20, 3: 16, 4: 14, 5: 12}
            run.font.size = Pt(sizes.get(level, 12))
            run.font.color.rgb = RGBColor.from_string('1E3A5F')
            para.space_before = Pt(12)
            para.space_after = Pt(6)
        
        elif elem_type == 'paragraph':
            add_formatted_paragraph(doc, element['text'])
        
        elif elem_type == 'table':
            table_data = parse_table(element['lines'])
            add_table_to_docx(doc, table_data)
        
        elif elem_type == 'bullet_list':
            for item in element['items']:
                para = doc.add_paragraph(style='List Bullet')
                parts = parse_inline_formatting(item)
                for part_text, formatting in parts:
                    add_formatted_run(para, part_text, formatting)
        
        elif elem_type == 'numbered_list':
            for item in element['items']:
                para = doc.add_paragraph(style='List Number')
                parts = parse_inline_formatting(item)
                for part_text, formatting in parts:
                    add_formatted_run(para, part_text, formatting)
        
        elif elem_type == 'blockquote':
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(1)
            run = para.add_run(element['text'])
            run.italic = True
            run.font.color.rgb = RGBColor(80, 80, 80)
        
        elif elem_type == 'code':
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(0.5)
            run = para.add_run(element['content'])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)

    # Add footer with generation info
    doc.add_page_break()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(f"Generated by Forward Thinking - Foresight\n{datetime.now().strftime('%B %d, %Y')}")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    footer_run.italic = True

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ============================================================================
# PDF EXPORT
# ============================================================================

def get_pdf_styles():
    """Create custom styles for PDF export."""
    styles = getSampleStyleSheet()
    
    # Custom heading styles
    styles.add(ParagraphStyle(
        name='Heading1Custom',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=12,
        spaceBefore=18,
        textColor=colors.HexColor('#1E3A5F'),
        fontName='Helvetica-Bold'
    ))
    
    styles.add(ParagraphStyle(
        name='Heading2Custom',
        parent=styles['Heading2'],
        fontSize=20,
        spaceAfter=10,
        spaceBefore=14,
        textColor=colors.HexColor('#2E5A7C'),
        fontName='Helvetica-Bold'
    ))
    
    styles.add(ParagraphStyle(
        name='Heading3Custom',
        parent=styles['Heading3'],
        fontSize=16,
        spaceAfter=8,
        spaceBefore=12,
        textColor=colors.HexColor('#3E7A9C'),
        fontName='Helvetica-Bold'
    ))
    
    styles.add(ParagraphStyle(
        name='Heading4Custom',
        parent=styles['Heading4'],
        fontSize=14,
        spaceAfter=6,
        spaceBefore=10,
        textColor=colors.HexColor('#4E8AAC'),
        fontName='Helvetica-Bold'
    ))
    
    styles.add(ParagraphStyle(
        name='Heading5Custom',
        parent=styles['Normal'],
        fontSize=12,
        spaceAfter=6,
        spaceBefore=8,
        textColor=colors.HexColor('#5E9ABC'),
        fontName='Helvetica-Bold'
    ))
    
    styles.add(ParagraphStyle(
        name='BodyCustom',
        parent=styles['Normal'],
        fontSize=11,
        leading=14,
        spaceAfter=8,
        alignment=TA_JUSTIFY
    ))
    
    styles.add(ParagraphStyle(
        name='BlockQuote',
        parent=styles['Normal'],
        fontSize=10,
        leading=13,
        leftIndent=30,
        rightIndent=20,
        spaceAfter=8,
        textColor=colors.HexColor('#505050'),
        fontName='Helvetica-Oblique'
    ))
    
    styles.add(ParagraphStyle(
        name='CodeBlock',
        parent=styles['Normal'],
        fontSize=9,
        leading=11,
        leftIndent=20,
        spaceAfter=8,
        fontName='Courier',
        backColor=colors.HexColor('#F5F5F5')
    ))
    
    styles.add(ParagraphStyle(
        name='TitleCustom',
        parent=styles['Title'],
        fontSize=28,
        spaceAfter=30,
        textColor=colors.HexColor('#1E3A5F'),
        alignment=TA_CENTER
    ))
    
    return styles


def convert_inline_to_reportlab(text: str) -> str:
    """Convert markdown inline formatting to ReportLab XML tags."""
    # Bold + Italic
    text = re.sub(r'\*\*\*(.+?)\*\*\*', r'<b><i>\1</i></b>', text)
    # Bold
    text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
    text = re.sub(r'__(.+?)__', r'<b>\1</b>', text)
    # Italic
    text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
    text = re.sub(r'_(.+?)_', r'<i>\1</i>', text)
    # Underline
    text = re.sub(r'<u>(.+?)</u>', r'<u>\1</u>', text)
    # Superscript
    text = re.sub(r'<sup>(.+?)</sup>', r'<super>\1</super>', text)
    text = re.sub(r'\^(.+?)\^', r'<super>\1</super>', text)
    # Subscript
    text = re.sub(r'<sub>(.+?)</sub>', r'<sub>\1</sub>', text)
    text = re.sub(r'~([^~]+?)~', r'<sub>\1</sub>', text)
    # Inline code
    text = re.sub(r'`(.+?)`', r'<font name="Courier" size="9">\1</font>', text)
    # Strikethrough (approximate with color)
    text = re.sub(r'~~(.+?)~~', r'<strike>\1</strike>', text)
    
    return text


def create_pdf_table(table_data: dict):
    """Create a formatted table for PDF."""
    headers = table_data['headers']
    rows = table_data['rows']
    
    if not headers:
        return None
    
    # Build table data
    data = []
    
    # Header row
    header_row = [Paragraph(f'<b>{h}</b>', getSampleStyleSheet()['Normal']) for h in headers]
    data.append(header_row)
    
    # Data rows
    for row in rows:
        data_row = [Paragraph(convert_inline_to_reportlab(cell), getSampleStyleSheet()['Normal']) for cell in row]
        data.append(data_row)
    
    # Create table
    col_widths = [None] * len(headers)  # Auto width
    table = Table(data, colWidths=col_widths)
    
    # Style the table
    style = TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#D9E2F3')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#1E3A5F')),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F9F9F9')]),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
        ('RIGHTPADDING', (0, 0), (-1, -1), 8),
    ])
    table.setStyle(style)
    
    return table


def export_to_pdf(content: str, title: str) -> bytes:
    """
    Export Markdown content as a richly formatted PDF document.
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=2*cm,
        leftMargin=2*cm,
        topMargin=2*cm,
        bottomMargin=2*cm
    )
    
    styles = get_pdf_styles()
    story = []
    
    # Add title
    story.append(Paragraph(title, styles['TitleCustom']))
    story.append(Spacer(1, 20))
    
    # Parse content
    elements = parse_markdown_content(content)
    
    for element in elements:
        elem_type = element.get('type')
        
        if elem_type == 'empty':
            story.append(Spacer(1, 6))
        
        elif elem_type == 'heading1':
            text = convert_inline_to_reportlab(element['text'])
            story.append(Paragraph(text, styles['Heading1Custom']))
        
        elif elem_type == 'heading2':
            text = convert_inline_to_reportlab(element['text'])
            story.append(Paragraph(text, styles['Heading2Custom']))
        
        elif elem_type == 'heading3':
            text = convert_inline_to_reportlab(element['text'])
            story.append(Paragraph(text, styles['Heading3Custom']))
        
        elif elem_type == 'heading4':
            text = convert_inline_to_reportlab(element['text'])
            story.append(Paragraph(text, styles['Heading4Custom']))
        
        elif elem_type == 'heading5':
            text = convert_inline_to_reportlab(element['text'])
            story.append(Paragraph(text, styles['Heading5Custom']))
        
        elif elem_type == 'paragraph':
            text = convert_inline_to_reportlab(element['text'])
            story.append(Paragraph(text, styles['BodyCustom']))
        
        elif elem_type == 'table':
            table_data = parse_table(element['lines'])
            table = create_pdf_table(table_data)
            if table:
                story.append(table)
                story.append(Spacer(1, 12))
        
        elif elem_type == 'bullet_list':
            bullet_items = []
            for item in element['items']:
                text = convert_inline_to_reportlab(item)
                bullet_items.append(ListItem(Paragraph(text, styles['BodyCustom'])))
            story.append(ListFlowable(bullet_items, bulletType='bullet', start='•'))
            story.append(Spacer(1, 6))
        
        elif elem_type == 'numbered_list':
            num_items = []
            for item in element['items']:
                text = convert_inline_to_reportlab(item)
                num_items.append(ListItem(Paragraph(text, styles['BodyCustom'])))
            story.append(ListFlowable(num_items, bulletType='1'))
            story.append(Spacer(1, 6))
        
        elif elem_type == 'blockquote':
            text = convert_inline_to_reportlab(element['text'])
            story.append(Paragraph(text, styles['BlockQuote']))
        
        elif elem_type == 'code':
            code_text = element['content'].replace('\n', '<br/>')
            story.append(Paragraph(code_text, styles['CodeBlock']))

    # Add footer with generation info
    from datetime import datetime
    story.append(Spacer(1, 30))
    footer_style = ParagraphStyle(
        'Footer',
        parent=styles['BodyCustom'],
        fontSize=9,
        textColor=colors.HexColor('#808080'),
        alignment=TA_CENTER,
        fontName='Helvetica-Oblique'
    )
    footer_text = f"Generated by Forward Thinking - Foresight<br/>{datetime.now().strftime('%B %d, %Y')}"
    story.append(Paragraph(footer_text, footer_style))

    # Build PDF
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


# ============================================================================
# MARKDOWN EXPORT (simple)
# ============================================================================

def export_to_markdown(content: str, title: str) -> bytes:
    """Export content as a Markdown file with metadata."""
    from datetime import datetime

    # Add metadata header
    metadata = f"""---
title: {title}
generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}
generator: Forward Thinking - Foresight
---

"""

    full_content = f"{metadata}# {title}\n\n{content}\n\n---\n*Generated by Forward Thinking - Foresight • {datetime.now().strftime('%B %d, %Y')}*"
    return full_content.encode('utf-8')


# ============================================================================
# PLANNING TABLE EXPORT
# ============================================================================

def export_planning_table_to_markdown(planning_data: dict) -> bytes:
    """Export the planning table as a Markdown file."""
    content = []
    
    overview = planning_data.get('book_overview', {})
    content.append(f"# Foresight Planning Table")
    content.append(f"\n## Book Overview\n")
    content.append(f"**Title:** {overview.get('title', 'N/A')}")
    content.append(f"**Scope:** {overview.get('scope', 'N/A')}")
    content.append(f"**Total Chapters:** {overview.get('total_chapters', 'N/A')}")
    content.append(f"**Disciplines:** {', '.join(overview.get('disciplines', []))}")
    
    content.append(f"\n## Planning Table\n")
    
    for chapter in planning_data.get('chapters', []):
        content.append(f"### {chapter.get('group_id', 'Unknown')}")
        content.append(f"**Chapters:** {', '.join(map(str, chapter.get('chapter_numbers', [])))}")
        content.append(f"**Titles:** {', '.join(chapter.get('chapter_titles', []))}")
        content.append(f"\n**Summary:** {chapter.get('content_summary', '')}")
        content.append(f"\n**Thematic Quadrants:** {', '.join(chapter.get('thematic_quadrants', []))}")
        content.append(f"\n**Foresight Task:**\n{chapter.get('foresight_task', '')}")
        content.append("\n---\n")
    
    if planning_data.get('implementation_notes'):
        content.append(f"## Implementation Notes\n")
        content.append(planning_data['implementation_notes'])
    
    return '\n'.join(content).encode('utf-8')


def export_planning_table_to_docx(planning_data: dict) -> bytes:
    """Export the planning table as a formatted DOCX file."""
    content = []
    
    overview = planning_data.get('book_overview', {})
    content.append(f"# Foresight Planning Table")
    content.append(f"\n## Book Overview\n")
    content.append(f"**Title:** {overview.get('title', 'N/A')}")
    content.append(f"**Scope:** {overview.get('scope', 'N/A')}")
    content.append(f"**Total Chapters:** {overview.get('total_chapters', 'N/A')}")
    content.append(f"**Disciplines:** {', '.join(overview.get('disciplines', []))}")
    
    for chapter in planning_data.get('chapters', []):
        content.append(f"\n## {chapter.get('group_id', 'Unknown')}")
        content.append(f"**Chapters:** {', '.join(map(str, chapter.get('chapter_numbers', [])))}")
        content.append(f"\n**Summary:** {chapter.get('content_summary', '')}")
        content.append(f"\n**Thematic Quadrants:** {', '.join(chapter.get('thematic_quadrants', []))}")
        content.append(f"\n**Foresight Task:**\n{chapter.get('foresight_task', '')}")
    
    markdown_content = '\n'.join(content)
    return export_to_docx(markdown_content, "Foresight Planning Table")


def export_planning_table_to_pdf(planning_data: dict) -> bytes:
    """Export the planning table as a formatted PDF file."""
    content = []
    
    overview = planning_data.get('book_overview', {})
    content.append(f"## Book Overview\n")
    content.append(f"**Title:** {overview.get('title', 'N/A')}")
    content.append(f"**Scope:** {overview.get('scope', 'N/A')}")
    content.append(f"**Total Chapters:** {overview.get('total_chapters', 'N/A')}")
    content.append(f"**Disciplines:** {', '.join(overview.get('disciplines', []))}")
    
    for chapter in planning_data.get('chapters', []):
        content.append(f"\n## {chapter.get('group_id', 'Unknown')}")
        content.append(f"**Chapters:** {', '.join(map(str, chapter.get('chapter_numbers', [])))}")
        content.append(f"\n**Summary:** {chapter.get('content_summary', '')}")
        content.append(f"\n**Thematic Quadrants:** {', '.join(chapter.get('thematic_quadrants', []))}")
        content.append(f"\n**Foresight Task:**\n{chapter.get('foresight_task', '')}")
    
    markdown_content = '\n'.join(content)
    return export_to_pdf(markdown_content, "Foresight Planning Table")
